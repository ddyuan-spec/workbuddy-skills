#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
html_to_docx.py  v2.2  ——  高保真 HTML -> Word(.docx) 转换器
v2.1 修复（2026-07-29）：
  - 修复中文乱码：每个 run 显式设置 w:eastAsia=Microsoft YaHei + post-process 补漏
  - 图片嵌入沿用 v1 已验证路径（paragraph.add_picture），修复 heading 内 img 丢失问题
  - 增加 post-process 验证与诊断输出
v2.2 修复（2026-07-29）：
  - 修复 HTML 注释泄漏为正文（bs4.Comment 是 NavigableString 子类，导致 <!-- --> 内容被当文本输出，
    表现为"标题重复"：注释中的章节名 + 正式 hN 标题同时出现）
  - 所有遍历 child 的节点（add_block / add_heading / unwrap_heading_blocks）统一跳过 Comment
用法：python html_to_docx.py <input.html> <output.docx> [--edge <msedge.exe>]
"""
import sys, os, re, io, subprocess, argparse
from bs4 import BeautifulSoup, NavigableString, Tag, Comment
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLOCK_TAGS = {"p","ul","ol","table","div","section","h1","h2","h3","h4","h5","img","svg"}

def find_edge():
    for p in [r"C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe",
              r"C:/Program Files/Microsoft/Edge/Application/msedge.exe"]:
        if os.path.exists(p): return p
    return "msedge"

# ── 核心：强制给每个 run 设置 CJK 字体 ──
def _cjk(run, font="Microsoft YaHei"):
    """在每个 run 上显式设置 eastAsia 字体，解决 WPS/Word 中文乱码。"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)

def set_cjk(style, font="Microsoft YaHei"):
    style.font.name = font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)

def shade(paragraph, fill="F8FAFC"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def add_hyperlink(paragraph, url, text):
    if not text: return
    part = paragraph.part
    rId = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), rId)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr'); rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'),'Hyperlink')
    rPr.append(rStyle); new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)

# ── 内联元素处理（每个 run 都走 _cjk）──
def add_inline(paragraph, node, bold=False):
    if isinstance(node, NavigableString):
        t = str(node)
        if t == "": return
        run = paragraph.add_run(t); run.bold = bold
        _cjk(run)
        return
    tag = node.name
    if tag in ("b","strong"):
        for ch in node.children: add_inline(paragraph, ch, bold=True)
    elif tag == "code":
        for ch in node.children:
            txt = str(ch) if isinstance(ch,NavigableString) else ch.get_text()
            run = paragraph.add_run(txt)
            run.font.name = "Consolas"; run.font.color.rgb = RGBColor(0xbe,0x12,0x3c)
            _cjk(run)
    elif tag == "a":
        add_hyperlink(paragraph, node.get("href",""), node.get_text())
    elif tag in ("span","em","i","font"):
        for ch in node.children: add_inline(paragraph, ch, bold=bold)
    elif tag == "br":
        paragraph.add_run().add_break()
    elif tag in ("img", "svg"):
        # 内联上下文中的图片/SVG：标记待处理，由调用方在 block 层处理
        # 这里仅记录 alt 文本作为 fallback
        alt = node.get("alt", "") if tag == "img" else f"[流程图]"
        if alt:
            run = paragraph.add_run(f"\n[图: {alt}]\n"); run.font.color.rgb = RGBColor(0x88,0x88,0x88); _cjk(run)
    else:
        for ch in node.children: add_inline(paragraph, ch, bold=bold)

def add_list(parent, el, ordered):
    for li in el.find_all("li", recursive=False):
        p = parent.add_paragraph(style="List Number" if ordered else "List Bullet")
        for ch in li.children:
            if isinstance(ch, NavigableString): add_inline(p, ch)
            elif ch.name in ("ul","ol"): add_list(parent, ch, ch.name=="ol")
            else: add_inline(p, ch)

# ── 图片嵌入（用 python-docx 原生 API，已验证可靠）──
def add_img(parent, el, width=5.8):
    src = el.get("src","")
    alt = el.get("alt","")
    if not src or src.startswith("http"):
        if alt:
            p = parent.add_paragraph(); r = p.add_run("[缺失图片] "+alt); r.font.color.rgb=RGBColor(0xdc,0x26,0x26); _cjk(r)
        return
    path = os.path.join(BASE, src)
    if not os.path.exists(path):
        if alt:
            p = parent.add_paragraph(); r = p.add_run("[缺失图片] "+alt); r.font.color.rgb=RGBColor(0xdc,0x26,0x26); _cjk(r)
        return
    p = parent.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    with open(path,"rb") as f: data = f.read()
    run = p.add_run()
    picture = run.add_picture(io.BytesIO(data), width=Inches(width))
    if alt:
        cap = parent.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x6b,0x72,0x80)
        _cjk(r)

def add_svg(parent, el, width=6.2):
    global SVG_COUNTER
    idx = SVG_COUNTER; SVG_COUNTER += 1
    png = SVG_PNGS[idx] if idx < len(SVG_PNGS) else None
    if png and os.path.exists(png):
        p = parent.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        with open(png,"rb") as f: data = f.read()
        run = p.add_run()
        run.add_picture(io.BytesIO(data), width=Inches(width))
    else:
        p = parent.add_paragraph(); r = p.add_run("[流程图渲染失败]"); r.font.color.rgb=RGBColor(0xdc,0x26,0x26); _cjk(r)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None: tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'999999')
        borders.append(e)
    tblPr.append(borders)

def add_table(parent, el):
    rows = el.find_all("tr")
    if not rows: return
    n_cols = max(len(r.find_all(["td","th"])) for r in rows)
    t = parent.add_table(rows=len(rows), cols=n_cols)
    try: t.style = "Table Grid"
    except Exception: pass
    set_table_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, r in enumerate(rows):
        cells = r.find_all(["td","th"])
        is_header = (r.find_parent("thead") is not None) or (cells and cells[0].name == "th")
        for ci in range(n_cols):
            cell = t.cell(ri, ci); cell.text = ""
            if ci >= len(cells): continue
            c = cells[ci]
            block_ch = [ch for ch in c.children if not isinstance(ch, NavigableString)]
            has_block = any(bc.name in ("p","ul","ol","div","table") for bc in block_ch)
            first = True
            if has_block:
                for ch in c.children:
                    if isinstance(ch, NavigableString):
                        if str(ch).strip():
                            para = cell.paragraphs[0] if first else cell.add_paragraph()
                            add_inline(para, ch); first = False
                    elif ch.name in ("p","div"):
                        para = cell.paragraphs[0] if first else cell.add_paragraph()
                        add_inline(para, ch); first = False
                    elif ch.name in ("ul","ol"):
                        add_list(cell, ch, ch.name=="ol"); first = False
                    else:
                        para = cell.paragraphs[0] if first else cell.add_paragraph()
                        add_inline(para, ch); first = False
            else:
                add_inline(cell.paragraphs[0], c)
            if is_header:
                for para in cell.paragraphs:
                    for run in para.runs: run.bold = True

# ── 预处理：解包标题内错误嵌套的块级子节点（修复「标题吞正文」）──
def unwrap_heading_blocks(soup):
    """html.parser 在 <h3> 内遇到 <h4>/<table> 时不会像浏览器那样隐式闭合标题，
    导致后续内容被解析成该 h3 的子节点，转换器又把子节点全文塞进标题段落，
    表现为「标题巨大 / 大小标题内容错乱重复」。
    此处把标题内错误嵌套的【直接块级子节点】提升到与标题同级，恢复正确层级。
    （源 HTML 在浏览器中显示正常，问题仅出在 html.parser 的解析差异。）"""
    HEADING = ['h1','h2','h3','h4','h5','h6']
    BLOCK = {'p','div','table','ul','ol','pre','blockquote','section',
             'h1','h2','h3','h4','h5','h6','figure'}
    for h in soup.find_all(HEADING):
        # 逆序插入，保证文档顺序
        for b in reversed([c for c in h.children
                           if not isinstance(c, (NavigableString, Comment)) and c.name in BLOCK]):
            h.insert_after(b)

# ── 标题处理（只取直接文本 + 直接行内子节点，块级已在预处理解包）──
def add_heading(parent, el, level):
    p = parent.add_paragraph()
    try:
        p.style = doc.styles["Title"] if level == 0 else doc.styles[f"Heading {level}"]
    except Exception:
        pass
    for child in el.children:
        if isinstance(child, Comment):
            continue
        if isinstance(child, NavigableString):
            t = str(child)
            if t.strip():
                run = p.add_run(t); _cjk(run)
        elif child.name == "img":
            add_img(parent, child)
        elif child.name == "svg":
            add_svg(parent, child)
        elif child.name in ("b","strong","span","em","i","font","code","a"):
            add_inline(p, child)
        # 其余块级子节点已在预处理阶段 unwrap 到标题同级，此处跳过

def div_has_block(el):
    for ch in el.children:
        if not isinstance(ch, NavigableString) and ch.name in BLOCK_TAGS:
            return True
    return False

def add_block(parent, el):
    for child in el.children:
        # 跳过 HTML 注释（bs4.Comment 是 NavigableString 子类，会被 isinstance 匹配）
        if isinstance(child, Comment):
            continue
        if isinstance(child, NavigableString):
            if str(child).strip():
                p = parent.add_paragraph(); add_inline(p, child)
            continue
        name = child.name
        if name == "h1": add_heading(parent, child, 0)
        elif name == "h2": add_heading(parent, child, 1)
        elif name == "h3": add_heading(parent, child, 2)
        elif name == "h4": add_heading(parent, child, 3)
        elif name == "h5": add_heading(parent, child, 4)
        elif name == "p":
            p = parent.add_paragraph(); add_inline(p, child)
        elif name in ("ul","ol"): add_list(parent, child, name=="ol")
        elif name == "table": add_table(parent, child)
        elif name == "img": add_img(parent, child)
        elif name == "svg": add_svg(parent, child)
        elif name in ("div","section"):
            cls = child.get("class", []) or []
            if div_has_block(child):
                add_block(parent, child)
            else:
                p = parent.add_paragraph()
                if "note" in cls: shade(p)
                if "pending" in cls: shade(p, "FFF7ED")
                add_inline(p, child)
        else:
            add_block(parent, child)

def render_svgs(soup, edge, tmp):
    nodes = soup.find_all("svg")
    out = []
    for i, svg in enumerate(nodes):
        vb = svg.get("viewBox", "0 0 1000 500").split()
        w, h = int(vb[2]), int(vb[3])
        s = str(svg)
        s = re.sub(r'style="[^"]*"', '', s)
        s = re.sub(r'<svg\b', f'<svg width="{w}" height="{h}"', s, count=1)
        wrap = f'<!DOCTYPE html><html><head><meta charset="utf-8"></head><body style="margin:0;padding:0;background:#fff;">{s}</body></html>'
        html_path = f"{tmp}/svg_{i}.html"
        png_path = f"{tmp}/svg_{i}.png"
        open(html_path,"w",encoding="utf-8").write(wrap)
        url = "file:///" + html_path.replace("\\","/")
        subprocess.run([edge,"--headless=new","--disable-gpu","--no-sandbox","--hide-scrollbars",
                        "--force-device-scale-factor=2","--window-size=%d,%d"%(w,h),
                        "--screenshot="+png_path, url], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        out.append(png_path)
    return out

# ── 后处理验证 ──
def post_process(doc):
    """遍历所有段落，补漏 CJK 字体 + 输出诊断信息。"""
    stats = {"total_runs": 0, "had_cjk": 0, "fixed": 0}
    for p in doc.paragraphs:
        for run in p.runs:
            stats["total_runs"] += 1
            rPr = run._r.find(qn('w:rPr'))
            has_ea = False
            if rPr is not None:
                rf = rPr.find(qn('w:rFonts'))
                if rf is not None and rf.get(qn('w:eastAsia')):
                    has_ea = True
            if has_ea:
                stats["had_cjk"] += 1
            else:
                _cjk(run)
                stats["fixed"] += 1
    print(f"  FONT: runs={stats['total_runs']} had_eastAsia={stats['had_cjk']} fixed={stats['fixed']}")
    print(f"  IMAGES: inline_shapes={len(doc.inline_shapes)} tables={len(doc.tables)} paras={len(doc.paragraphs)}")
    # 标题健康检查：检测「标题吞正文 / 标题重复错乱」
    # 成因：源 HTML 中 hN 标签开闭名不匹配（如 <h4>...</h5>、<h3>...</h4>），
    #       html.parser 不会隐式闭合，导致后续整段被吞进该标题段落。
    sus = 0
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') or p.style.name == 'Title':
            t = p.text
            if '\n' in t or len(t) > 60:
                sus += 1
                print(f"  ⚠️ HEADING HEALTH: [{p.style.name}] 疑似吞正文(len={len(t)}): {t[:50]!r}")
    print(f"  HEADING HEALTH: suspect={sus}")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("input"); ap.add_argument("output")
    ap.add_argument("--edge", default=None)
    args = ap.parse_args()
    SRC = os.path.abspath(args.input); BASE = os.path.dirname(SRC)
    OUT = os.path.abspath(args.output)
    EDGE = args.edge or find_edge()
    TMP = os.path.join(os.path.dirname(OUT), "_docx_tmp")
    os.makedirs(TMP, exist_ok=True)

    raw = open(SRC, encoding="utf-8").read()
    soup = BeautifulSoup(raw, "html.parser")

    SVG_PNGS = render_svgs(soup, EDGE, TMP)
    SVG_COUNTER = 0

    doc = Document()
    set_cjk(doc.styles["Normal"]); doc.styles["Normal"].font.size = Pt(10.5)
    pf = doc.styles["Normal"].paragraph_format
    pf.line_spacing = 1.5
    for lvl in (0,1,2,3,4):
        try:
            s = doc.styles["Title"] if lvl==0 else doc.styles[f"Heading {lvl}"]
            set_cjk(s)
            sp = s.paragraph_format
            sp.space_before = Pt(12 if lvl > 0 else 18)
            sp.space_after = Pt(6)
        except Exception: pass

    unwrap_heading_blocks(soup)

    body = soup.body or soup
    add_block(doc, body)
    
    post_process(doc)
    
    doc.save(OUT)
    print(f"OK saved {OUT}")
